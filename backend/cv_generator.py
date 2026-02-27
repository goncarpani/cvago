"""
Tool 2 — CV Generator: JSON (perfil) + Job Description → CV personalizado en PDF y DOCX.
El LLM reordena/reframe según JD respetando constraints.cannotModify. No inventa datos.
"""
import json
import os
from collections import OrderedDict
from pathlib import Path
from datetime import datetime

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import subprocess

OUTPUT_DIR = Path(__file__).resolve().parent / "generated_cvs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Colors matching the reference CV template
CLR_DARK = RGBColor(0x40, 0x40, 0x40)
CLR_BODY = RGBColor(0x59, 0x59, 0x59)
CLR_LINK = RGBColor(0x11, 0x55, 0xCC)

ROLE_INDENT = Cm(0.6)
DESC_INDENT = Cm(0.8)

CANNOT_MODIFY = [
    "dates (start, end de cada experiencia)",
    "companies (nombres de empresas)",
    "officialTitle (título oficial del puesto)",
    "technologies (lista de tecnologías)",
    "educationDegrees (títulos y datos de educación)",
]
CAN_REFRAME = [
    "achievements / bullets / facts",
    "summary / headline",
    "bulletOrdering (orden de los logros)",
    "skillHighlighting (qué skills destacar)",
    "capabilityEmphasis",
    "headline",
]

SYSTEM_PROMPT = f"""Generás CVs adaptados a una job description. Recibís un perfil en JSON y una JD.

Reglas INVIOLABLES (no negociables):
- NUNCA modifiques: {json.dumps(CANNOT_MODIFY, ensure_ascii=False)}
- No podés incluir información que no esté en el JSON del perfil. Nada inventado.
  - No inventes logros, métricas, tecnologías ni responsabilidades.
  - Cada frase debe poder trazarse a algún raw, fact, capability o technology del perfil.
- Respetá narrative.avoidFraming: no enfatices esos aspectos en el CV generado.

Idioma:
- El backend te indica explícitamente el idioma (español o inglés) en una instrucción aparte.
- NO infieras el idioma a partir de la JD: usá SIEMPRE el idioma indicado.

Output: respondé ÚNICAMENTE con un JSON válido con esta forma (sin markdown, sin explicaciones):
{{
  "headline": "Resumen corto alineado con la JD (máx. ~300 caracteres, idealmente hasta 3 líneas en el documento)",
  "experience": [
    {{
      "company": "(igual al perfil)",
      "officialTitle": "(igual al perfil)",
      "start": "YYYY-MM",
      "end": "YYYY-MM o present",
      "location": "",
      "description": "párrafo corrido con la información MÁS relevante de este rol para la JD"
    }}
  ],
  "education": [
    {{ "degree": "", "institution": "", "year": "", "location": "", "note": "" }}
  ],
  "skills_technical": "lista o string de skills técnicas relevantes para la JD (no todo el inventario del perfil)",
  "skills_soft": "lista o string de soft skills relevantes para la JD (en particular, liderazgo solo si la JD lo pide o lo sugiere)",
  "languages": "idiomas y niveles relevantes para la JD"
}}

Orden y omisión de roles:
- Ordená SIEMPRE experience en orden cronológico inverso estricto: el rol más reciente primero, el más antiguo al final.
- La experiencia actual (end == "present") DEBE aparecer siempre primera en el array experience.
- Los roles recientes SIEMPRE se incluyen, aunque no sean los más relevantes para la JD. Si hace falta espacio, comprimí su descripción, pero no los omitas.
- Solo podés omitir experiencias si el texto total del CV (todas las secciones) no entra en aproximadamente 3.800 caracteres.
- Si necesitás omitir, SIEMPRE empezá por las experiencias más antiguas (las de end más viejo) y nunca omitas la experiencia actual ni los roles recientes.
- No hay un límite fijo de cantidad de roles: usá tantos como entren respetando las reglas anteriores.

Formato de descripción por rol:
- NO uses bullets ni listas para las experiencias.
- Para cada experiencia incluída, generá un único campo "description" con un párrafo corrido.
- Ese párrafo debe combinar y reformular los facts y capabilities más relevantes de ese rol para esta JD, usando el vocabulario de la oferta.
- Roles más relevantes para la JD pueden tener descripciones más largas; roles menos relevantes deben comprimirse más.

Control de longitud global:
- El único constraint de espacio es la longitud total aproximada del documento: apuntá a ~3.800 caracteres en total para que quepa en una página.
- NO apliques límites fijos de líneas ni de caracteres por rol o por sección: distribuí el espacio según la relevancia de cada sección y cada rol para esta JD.
- Si necesitás comprimir, empezá por reducir detalle en roles antiguos o menos relevantes, manteniendo siempre la claridad.

Skills:
- La sección de skills NO es un inventario completo del perfil.
- Skills técnicas: listá ÚNICAMENTE tecnologías y herramientas directamente relevantes para la JD — mencionadas explícitamente en la oferta o claramente relacionadas con las responsabilidades descriptas. Si una tecnología del perfil no tiene relación con ninguna responsabilidad de la JD, no la incluyas aunque sea la más reciente o destacada del perfil.
- La selección de skills técnicas debe reflejar lo que el hiring manager de esa posición considera relevante. Para determinarlo, analizá el tipo de rol y las responsabilidades de la JD: si el rol es de operaciones, gestión o análisis, priorizá herramientas de análisis y reporting; si el rol es de ingeniería o desarrollo, priorizá herramientas de implementación y arquitectura; si el rol es híbrido, incluí lo más relevante de ambas categorías. En todos los casos, la regla principal aplica: solo incluir lo que tenga relación directa con la JD.
- Skills blandas y de liderazgo: cuando la JD menciona explícitamente stakeholder management, cross-functional collaboration, process improvement, communication o gestión de equipos, las soft skills son tan o más importantes que las técnicas y deben aparecer primero o con igual peso. En esos casos priorizalas; si no, mantené esa parte corta o omitila en favor de contenido más relevante.
- No hay mínimo ni máximo fijo de skills: priorizá relevancia para la JD y legibilidad.

Education:
- Incluí siempre la información completa del perfil: degree, institution, location, year, y cualquier nota (ej: "Coursework completed, thesis pending").
- "location" y "note" son campos opcionales — incluílos si el perfil los tiene.

Summary / headline:
- "headline" debe ser un resumen corto alineado con la JD, no un resumen genérico del perfil.
- La headline NUNCA debe redefinir la identidad profesional del candidato. Debe reflejar quién es según su trayectoria real (por ejemplo, narrative.headline y narrative.coreIdentity del perfil), adaptando el énfasis a la JD sin cambiar la identidad de base. Un líder de analytics no puede aparecer como software engineer, y un ingeniero de infraestructura no puede aparecer como analista de negocio.
- Máximo aproximado: 300 caracteres (lo que equivale a unas 2–3 líneas en el documento).
- Explicá brevemente por qué este candidato es relevante para ESTA posición específica (rol, seniority, contexto).

Recordatorio importante:
- Todo el contenido generado (headline, descriptions, skills, languages) debe estar en el idioma indicado por el backend.
- Empezá tu respuesta directamente con {{ y terminá con }} para que sea un JSON puro sin texto extra.
"""

LANGUAGE_LABELS = {
    "es": {"experience": "EXPERIENCIA", "education": "EDUCACIÓN", "skills": "SKILLS"},
    "en": {"experience": "EXPERIENCE", "education": "EDUCATION", "skills": "SKILLS"},
}


def generate_cv_content(profile: dict, jd_text: str, language: str = "es") -> dict:
    """
    Llama al LLM con perfil + JD y devuelve la estructura lista para renderizar.
    language: "es" (español) o "en" (inglés). Todo el contenido generado va en ese idioma.
    """
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY no configurada. Necesaria para el CV Generator.")

    lang = "en" if language == "en" else "es"
    lang_instruction = (
        "Generate the entire CV in English. All text (headline, experience descriptions, skills_technical, skills_soft, languages) must be in English."
        if lang == "en"
        else "Generá todo el CV en español. Todo el texto (headline, descripciones de experiencia, skills_technical, skills_soft, languages) debe estar en español."
    )

    client = OpenAI(api_key=api_key)
    model = os.environ.get("OPENAI_MODEL", "gpt-5.2")
    jd_chunk = jd_text[:15000].strip() if jd_text else ""
    profile_str = json.dumps(profile, ensure_ascii=False, indent=0)[:25000]

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"Perfil (JSON):\n{profile_str}\n\n---\nJob Description:\n{jd_chunk}\n\n---\n{lang_instruction}\n\nDevolvé el JSON del CV adaptado.",
            },
        ],
        max_completion_tokens=6000,
        temperature=0.2,
    )
    raw = (response.choices[0].message.content or "").strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
    return json.loads(raw)


def _escape(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_MONTHS_ES = [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]
_MONTHS_EN = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def _format_date(value: str, language: str) -> str:
    """Convierte 'YYYY-MM' a 'Month YYYY' según idioma. 'present' → 'Presente'/'Present'."""
    if not value or not isinstance(value, str):
        return value or ""
    v = value.strip().lower()
    if v == "present":
        return "Presente" if language == "es" else "Present"
    if len(v) == 7 and v[4] == "-":
        try:
            y, m = int(v[:4]), int(v[5:7])
            if 1 <= m <= 12:
                months = _MONTHS_ES if language == "es" else _MONTHS_EN
                return f"{months[m - 1]} {y}"
        except ValueError:
            pass
    return value


def _extract_year(date_str: str) -> str:
    """Extract just the year from a date string like '2023-06' or 'present'."""
    if not date_str:
        return ""
    v = date_str.strip().lower()
    if v == "present":
        return "Present"
    if len(v) >= 4:
        return v[:4]
    return date_str


def _group_experiences(experiences: list) -> list:
    """Group flat experience entries by company, preserving order.
    Returns list of (company, year_range, roles) tuples."""
    groups = OrderedDict()
    for exp in experiences:
        company = exp.get("company", "Unknown")
        if company not in groups:
            groups[company] = []
        groups[company].append(exp)

    result = []
    for company, roles in groups.items():
        starts = [r.get("start", "") for r in roles if r.get("start")]
        ends = [r.get("end", "") for r in roles if r.get("end")]
        start_year = _extract_year(min(starts)) if starts else ""
        end_raw = max(ends, key=lambda x: "9999" if x.strip().lower() == "present" else x) if ends else ""
        end_year = "Present" if end_raw.strip().lower() == "present" else _extract_year(end_raw)
        year_range = f"{start_year} - {end_year}" if start_year and end_year else ""
        result.append((company, year_range, roles))
    return result


def _get_description(exp: dict) -> str:
    """Get description from experience, with backwards compat for bullets."""
    description = exp.get("description")
    if not description and isinstance(exp.get("bullets"), list):
        description = " ".join(str(b) for b in exp["bullets"] if b)
    return description or ""


def _add_bottom_border(paragraph, color="404040", size="6"):
    """Add a horizontal line below a paragraph (section separator)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ─────────────────────────────── DOCX ──────────────────────────────

def _add_run(paragraph, text, font_name="Tahoma", size_pt=10, bold=True, color=None):
    """Helper to add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def write_docx(cv_content: dict, profile: dict, output_path: Path, language: str = "es") -> None:
    """Escribe el CV en DOCX replicando el estilo del template de referencia."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(30.48)
    section.left_margin = Cm(1.31)
    section.right_margin = Cm(1.21)
    section.top_margin = Cm(1.36)
    section.bottom_margin = Cm(1.17)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Tahoma"
    normal_style.font.size = Pt(10)

    personal = profile.get("personal", {})
    name = f"{personal.get('firstName', '')} {personal.get('lastName', '')}".strip()

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    _add_run(p, name, size_pt=17, color=CLR_DARK)

    # Contact line
    contact_parts = list(filter(None, [personal.get("location"), personal.get("phone"), personal.get("email")]))
    if contact_parts:
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c.paragraph_format.space_before = Pt(3)
        p_c.paragraph_format.space_after = Pt(0)
        p_c.paragraph_format.line_spacing = 1.0
        for i, part in enumerate(contact_parts):
            if i > 0:
                _add_run(p_c, " | ", size_pt=10, color=CLR_BODY)
            is_email = "@" in part
            _add_run(p_c, part, size_pt=10, color=CLR_LINK if is_email else CLR_BODY)

    # Links line
    links = personal.get("links", {})
    link_values = [v for v in links.values() if v]
    if link_values:
        p_l = doc.add_paragraph()
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_l.paragraph_format.space_before = Pt(1)
        p_l.paragraph_format.space_after = Pt(0)
        p_l.paragraph_format.line_spacing = 1.0
        for i, link in enumerate(link_values):
            if i > 0:
                _add_run(p_l, " | ", size_pt=10, color=CLR_BODY)
            _add_run(p_l, link, size_pt=10, color=CLR_LINK)

    # Headline / summary
    if cv_content.get("headline"):
        p_h = doc.add_paragraph()
        p_h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_h.paragraph_format.space_before = Pt(3)
        p_h.paragraph_format.space_after = Pt(0)
        p_h.paragraph_format.line_spacing = 1.07
        _add_run(p_h, cv_content["headline"], size_pt=10, color=CLR_BODY)

    labels = LANGUAGE_LABELS.get(language, LANGUAGE_LABELS["es"])

    # ── EXPERIENCE ──
    p_exp_h = doc.add_paragraph()
    p_exp_h.paragraph_format.space_before = Pt(8)
    p_exp_h.paragraph_format.space_after = Pt(0)
    p_exp_h.paragraph_format.line_spacing = 1.0
    _add_run(p_exp_h, labels["experience"], size_pt=12, color=CLR_DARK)
    _add_bottom_border(p_exp_h)

    for company, year_range, roles in _group_experiences(cv_content.get("experience", [])):
        # Company heading
        p_comp = doc.add_paragraph()
        p_comp.paragraph_format.space_before = Pt(6)
        p_comp.paragraph_format.space_after = Pt(0)
        p_comp.paragraph_format.line_spacing = 1.07
        company_text = f"{company} – {year_range}" if year_range else company
        _add_run(p_comp, company_text, size_pt=11, color=CLR_DARK)

        for role in roles:
            start_fmt = _format_date(role.get("start", ""), language)
            end_fmt = _format_date(role.get("end", ""), language)

            # Role line: ● Title | Date range:  (indented)
            p_role = doc.add_paragraph()
            p_role.paragraph_format.space_before = Pt(2)
            p_role.paragraph_format.space_after = Pt(0)
            p_role.paragraph_format.line_spacing = 1.0
            p_role.paragraph_format.left_indent = ROLE_INDENT
            _add_run(p_role, "● ", font_name="Arial", size_pt=10, color=CLR_BODY)
            _add_run(p_role, f"{role.get('officialTitle', '')} | {start_fmt} – {end_fmt}:", size_pt=10, color=CLR_BODY)

            # Description (indented)
            desc = _get_description(role)
            if desc:
                p_desc = doc.add_paragraph()
                p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_desc.paragraph_format.space_before = Pt(0.5)
                p_desc.paragraph_format.space_after = Pt(0)
                p_desc.paragraph_format.line_spacing = 1.07
                p_desc.paragraph_format.left_indent = DESC_INDENT
                _add_run(p_desc, desc, size_pt=10, color=CLR_BODY)

    # ── EDUCATION ──
    p_edu_h = doc.add_paragraph()
    p_edu_h.paragraph_format.space_before = Pt(8)
    p_edu_h.paragraph_format.space_after = Pt(0)
    p_edu_h.paragraph_format.line_spacing = 1.0
    _add_run(p_edu_h, labels["education"], size_pt=11, color=CLR_DARK)
    _add_bottom_border(p_edu_h)

    for ed in cv_content.get("education", []):
        p_ed = doc.add_paragraph()
        p_ed.paragraph_format.space_before = Pt(2)
        p_ed.paragraph_format.space_after = Pt(0)
        p_ed.paragraph_format.line_spacing = 1.0
        degree = ed.get("degree", "")
        institution = ed.get("institution", "")
        location = ed.get("location", "")
        year = ed.get("year", "")
        note = ed.get("note", "")
        _add_run(p_ed, degree, size_pt=11, color=CLR_DARK)
        if institution:
            loc_str = f" - {location}" if location else ""
            _add_run(p_ed, f", {institution}{loc_str}; {year}", size_pt=10, color=CLR_DARK)
        elif year:
            _add_run(p_ed, f" ({year})", size_pt=10, color=CLR_DARK)
        if note:
            _add_run(p_ed, f" ({note})", size_pt=10, color=CLR_DARK)

    # ── SKILLS (includes Languages) ──
    p_sk_h = doc.add_paragraph()
    p_sk_h.paragraph_format.space_before = Pt(8)
    p_sk_h.paragraph_format.space_after = Pt(0)
    p_sk_h.paragraph_format.line_spacing = 1.0
    _add_run(p_sk_h, labels["skills"], size_pt=11, color=CLR_DARK)
    _add_bottom_border(p_sk_h)

    soft = cv_content.get("skills_soft") or ""
    if isinstance(soft, list):
        soft = ", ".join(soft)
    tech = cv_content.get("skills_technical") or ""
    if isinstance(tech, list):
        tech = ", ".join(tech)
    lang_text = str(cv_content.get("languages", "")) if cv_content.get("languages") else ""

    if soft:
        p_soft = doc.add_paragraph()
        p_soft.paragraph_format.space_before = Pt(2)
        p_soft.paragraph_format.space_after = Pt(0)
        p_soft.paragraph_format.line_spacing = 1.07
        soft_label = "People & Leadership" if language == "en" else "Personas & Liderazgo"
        _add_run(p_soft, f"{soft_label}: ", size_pt=9, color=CLR_BODY)
        _add_run(p_soft, soft, size_pt=9, color=CLR_BODY)

    if tech:
        p_tech = doc.add_paragraph()
        p_tech.paragraph_format.space_before = Pt(2)
        p_tech.paragraph_format.space_after = Pt(0)
        p_tech.paragraph_format.line_spacing = 1.0
        tech_label = "Technical" if language == "en" else "Técnicas"
        _add_run(p_tech, f"{tech_label}: ", size_pt=9, color=CLR_BODY)
        _add_run(p_tech, tech, size_pt=9, color=CLR_BODY)

    if lang_text:
        p_lang = doc.add_paragraph()
        p_lang.paragraph_format.space_before = Pt(2)
        p_lang.paragraph_format.space_after = Pt(0)
        p_lang.paragraph_format.line_spacing = 1.0
        lang_label = "Languages" if language == "en" else "Idiomas"
        _add_run(p_lang, f"{lang_label}: ", size_pt=9, color=CLR_BODY)
        _add_run(p_lang, lang_text, size_pt=9, color=CLR_BODY)

    doc.save(str(output_path))


def _docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert DOCX to PDF by calling docx2pdf in a subprocess (avoids COM/DLL issues inside uvicorn)."""
    import sys
    result = subprocess.run(
        [sys.executable, "-c",
         f"from docx2pdf import convert; convert(r'{docx_path}', r'{pdf_path}')"],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"docx2pdf failed: {result.stderr.strip()}")
    if not pdf_path.exists():
        raise RuntimeError("PDF was not created by docx2pdf")


def generate_cv_pdf_and_docx(
    profile: dict, jd_text: str, language: str = "es", base_name: str | None = None
) -> tuple[str, str]:
    """
    Genera el CV adaptado: DOCX via python-docx, luego convierte a PDF via Word (docx2pdf).
    Devuelve (pdf_filename, docx_filename).
    """
    lang = "en" if language == "en" else "es"
    cv_content = generate_cv_content(profile, jd_text, language=lang)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = base_name or f"cv_{stamp}"
    docx_path = OUTPUT_DIR / f"{base}.docx"
    pdf_path = OUTPUT_DIR / f"{base}.pdf"
    write_docx(cv_content, profile, docx_path, language=lang)
    _docx_to_pdf(docx_path, pdf_path)
    return f"{base}.pdf", f"{base}.docx"
